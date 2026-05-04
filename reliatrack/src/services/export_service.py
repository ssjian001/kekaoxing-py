"""导出服务 — Excel (openpyxl) + PDF (reportlab) + Word (python-docx) 导出。"""

from __future__ import annotations

import logging
import os
import sys
from datetime import datetime
from pathlib import Path

from src.constants import (
    TASK_STATUS_LABELS,
    ISSUE_STATUS_LABELS,
    SAMPLE_STATUS_LABELS,
)
from src.models.test_plan import TestPlan, TestTask, TestResult
from src.models.issue import Issue, FARecord, CAPARecord
from src.models.sample import Sample

logger = logging.getLogger(__name__)


class ExportService:
    """导出服务 — 将数据导出为 Excel / PDF / Word 文件。"""

    # ── 跨平台 CJK 字体名（供 openpyxl / python-docx 使用）──
    _CJK_FONT: str | None = None

    @classmethod
    def get_cjk_font(cls) -> str:
        """返回当前平台可用的 CJK 字体名（单个字体，非 fallback 链）。

        优先级：Microsoft YaHei > PingFang SC > Noto Sans CJK SC > sans-serif。
        缓存到类变量，只检测一次。
        """
        if cls._CJK_FONT is not None:
            return cls._CJK_FONT
        import platform
        system = platform.system()
        if system == "Windows":
            cls._CJK_FONT = "Microsoft YaHei"
        elif system == "Darwin":
            cls._CJK_FONT = "PingFang SC"
        else:
            # Linux: 尝试检测已安装的 CJK 字体
            import subprocess
            for candidate in ("Noto Sans CJK SC", "WenQuanYi Micro Hei", "Droid Sans Fallback"):
                try:
                    r = subprocess.run(
                        ["fc-match", "-f", "%{family}", candidate],
                        capture_output=True, text=True, timeout=3,
                    )
                    if r.returncode == 0 and r.stdout.strip():
                        cls._CJK_FONT = candidate
                        break
                except Exception:
                    continue
            if cls._CJK_FONT is None:
                cls._CJK_FONT = "Noto Sans CJK SC"  # 合理的默认值
        return cls._CJK_FONT

    # ── 类别中文映射 ──
    CATEGORY_MAP = {
        "环境试验": "env",
        "env": "环境试验",
        "机械试验": "mech",
        "mech": "机械试验",
        "表面处理": "surf",
        "surf": "表面处理",
        "包装": "pack",
        "pack": "包装",
        "其他": "other",
        "other": "其他",
    }

    # Merged status → Chinese label map (task + issue + sample)
    STATUS_MAP: dict[str, str] = {
        **TASK_STATUS_LABELS,
        **ISSUE_STATUS_LABELS,
        **SAMPLE_STATUS_LABELS,
    }

    def __init__(self, output_dir: str = "exports") -> None:
        self._output_dir = Path(output_dir)

    def _ensure_dir(self) -> Path:
        self._output_dir.mkdir(parents=True, exist_ok=True)
        return self._output_dir

    # ── Excel 共享辅助 ──

    def _excel_styles(self, fill_color: str = "2B579A"):
        """返回 Excel 通用样式字典（字体、填充、对齐、边框）。"""
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

        _f = self.get_cjk_font()
        return {
            "header_font": Font(name=_f, size=11, bold=True, color="FFFFFF"),
            "header_fill": PatternFill(start_color=fill_color, end_color=fill_color, fill_type="solid"),
            "cell_font": Font(name=_f, size=10),
            "title_font": lambda c: Font(name=_f, size=14, bold=True, color=c),
            "sub_font": Font(name=_f, size=9, color="666666"),
            "center": Alignment(horizontal="center", vertical="center"),
            "thin_border": Border(
                left=Side(style="thin"), right=Side(style="thin"),
                top=Side(style="thin"), bottom=Side(style="thin"),
            ),
        }

    def _excel_write_title_block(
        self, ws, title: str, merge_range: str, subtitle: str, sub_range: str,
        title_color: str, styles: dict,
    ) -> None:
        """写入 Excel 标题行 + 副标题行。"""
        from openpyxl.styles import Alignment

        ws.merge_cells(merge_range)
        title_cell = ws[merge_range.split(":")[0]]
        title_cell.value = title
        title_cell.font = styles["title_font"](title_color)
        title_cell.alignment = Alignment(horizontal="center")
        ws.row_dimensions[1].height = 30

        ws.merge_cells(sub_range)
        sub = ws[sub_range.split(":")[0]]
        sub.value = subtitle
        sub.font = styles["sub_font"]
        sub.alignment = Alignment(horizontal="center")

    def _excel_write_headers(
        self, ws, row: int, headers: list[str], styles: dict,
    ) -> None:
        """写入 Excel 表头行。"""
        for col, h in enumerate(headers, 1):
            cell = ws.cell(row=row, column=col, value=h)
            cell.font = styles["header_font"]
            cell.fill = styles["header_fill"]
            cell.alignment = styles["center"]
            cell.border = styles["thin_border"]

    def _excel_write_row(
        self, ws, row: int, values: list, styles: dict, alignment=None,
    ) -> None:
        """写入 Excel 单行数据。"""
        align = alignment or styles["center"]
        for col, val in enumerate(values, 1):
            cell = ws.cell(row=row, column=col, value=val)
            cell.font = styles["cell_font"]
            cell.alignment = align
            cell.border = styles["thin_border"]

    def _excel_save(self, wb, filepath: str | None, filename: str) -> str:
        """保存 Excel 工作簿，返回绝对路径。"""
        out = filepath or str(self._ensure_dir() / filename)
        wb.save(out)
        return os.path.abspath(out)

    # ── Excel 导出 ──────────────────────────────────────────────

    def export_tasks_excel(
        self,
        plan: TestPlan,
        tasks: list[TestTask],
        filepath: str | None = None,
    ) -> str:
        """导出测试任务列表为 Excel。

        Returns:
            导出文件的绝对路径。
        """
        from openpyxl import Workbook

        s = self._excel_styles("2B579A")
        wb = Workbook()
        ws = wb.active
        ws.title = "测试任务"

        self._excel_write_title_block(
            ws, f"测试计划: {plan.name}", "A1:I1",
            f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}  |  测试标准: {plan.test_standard or '—'}",
            "A2:I2", "2B579A", s,
        )

        headers = ["#", "名称", "类别", "工期(天)", "开始天", "进度", "状态", "优先级", "环境条件"]
        self._excel_write_headers(ws, 4, headers, s)

        for row_idx, task in enumerate(tasks, 5):
            values = [
                row_idx - 4,
                task.name,
                self.CATEGORY_MAP.get(task.category, task.category),
                task.duration,
                task.start_day,
                f"{task.progress:.0f}%",
                self.STATUS_MAP.get(task.status, task.status),
                task.priority,
                task.environment,
            ]
            self._excel_write_row(ws, row_idx, values, s)

        widths = [5, 30, 10, 10, 10, 8, 10, 8, 25]
        for i, w in enumerate(widths, 1):
            ws.column_dimensions[chr(64 + i)].width = w

        return self._excel_save(wb, filepath, f"测试任务_{plan.name}_{datetime.now():%Y%m%d_%H%M}.xlsx")

    def export_issues_excel(
        self,
        issues: list[Issue],
        fa_map: dict[int, list[FARecord]] | None = None,
        filepath: str | None = None,
    ) -> str:
        """导出 Issue 列表为 Excel。

        Args:
            fa_map: {issue_id: [FARecord, ...]} 可选。
        """
        from openpyxl import Workbook
        from openpyxl.styles import Alignment

        s = self._excel_styles("C0504D")
        wrap = Alignment(horizontal="left", vertical="center", wrap_text=True)

        wb = Workbook()
        ws = wb.active
        ws.title = "Issue 追踪"

        self._excel_write_title_block(
            ws, "Issue 追踪报告", "A1:H1",
            f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}  |  共 {len(issues)} 个 Issue",
            "A2:H2", "C0504D", s,
        )

        headers = ["ID", "标题", "严重度", "状态", "优先级", "失效模式", "根因分析", "FA 步骤数"]
        self._excel_write_headers(ws, 4, headers, s)

        for row_idx, issue in enumerate(issues, 5):
            fa_count = len(fa_map.get(issue.id, [])) if fa_map and issue.id is not None else 0
            values = [
                issue.id,
                issue.title,
                issue.severity,
                self.STATUS_MAP.get(issue.status, issue.status),
                issue.priority,
                issue.failure_mode or "",
                (issue.root_cause or "")[:100],
                fa_count,
            ]
            for col, val in enumerate(values, 1):
                cell = ws.cell(row=row_idx, column=col, value=val)
                cell.font = s["cell_font"]
                cell.alignment = wrap if col in (2, 6, 7) else s["center"]
                cell.border = s["thin_border"]

        widths = [5, 25, 10, 10, 8, 15, 35, 10]
        for i, w in enumerate(widths, 1):
            ws.column_dimensions[chr(64 + i)].width = w

        return self._excel_save(wb, filepath, f"Issue追踪_{datetime.now():%Y%m%d_%H%M}.xlsx")

    def export_samples_excel(
        self,
        samples: list[Sample],
        filepath: str | None = None,
    ) -> str:
        """导出样品台账为 Excel。"""
        from openpyxl import Workbook

        s = self._excel_styles("4F81BD")

        wb = Workbook()
        ws = wb.active
        ws.title = "样品台账"

        self._excel_write_title_block(
            ws, "样品台账", "A1:F1",
            f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}  |  共 {len(samples)} 个样品",
            "A2:F2", "4F81BD", s,
        )

        headers = ["ID", "SN", "批次号", "规格型号", "状态", "存放位置"]
        self._excel_write_headers(ws, 4, headers, s)

        for row_idx, sample in enumerate(samples, 5):
            values = [
                sample.id,
                sample.sn,
                sample.batch_no or "",
                sample.spec or "",
                self.STATUS_MAP.get(sample.status, sample.status),
                sample.location or "",
            ]
            self._excel_write_row(ws, row_idx, values, s)

        widths = [5, 20, 15, 20, 10, 20]
        for i, w in enumerate(widths, 1):
            ws.column_dimensions[chr(64 + i)].width = w

        return self._excel_save(wb, filepath, f"样品台账_{datetime.now():%Y%m%d_%H%M}.xlsx")

    # ── PDF 导出 ──────────────────────────────────────────────

    @staticmethod
    def _find_cjk_font() -> tuple[str, str, int | None, int | None]:
        """跨平台查找可用的中文字体。

        Returns:
            (regular_path, bold_path, regular_subfont, bold_subfont)
            subfont 为 None 表示纯 TTF 文件，无需指定索引。
        """
        if sys.platform == "win32":
            windir = os.environ.get("WINDIR", r"C:\Windows")
            fd = os.path.join(windir, "Fonts")
            candidates: list[tuple[str, str, int | None, int | None]] = [
                # 微软雅黑: msyh.ttc(reg) + msyhbd.ttc(bold)
                (os.path.join(fd, "msyh.ttc"), os.path.join(fd, "msyhbd.ttc"), 0, 0),
                # msyh.ttc 单文件（部分版本 bold 在同一 ttc 内）
                (os.path.join(fd, "msyh.ttc"), os.path.join(fd, "msyh.ttc"), 0, 1),
                # 黑体
                (os.path.join(fd, "simhei.ttf"), os.path.join(fd, "simhei.ttf"), None, None),
            ]
        elif sys.platform == "darwin":
            candidates = [
                ("/System/Library/Fonts/PingFang.ttc",
                 "/System/Library/Fonts/PingFang.ttc", 1, 3),
            ]
        else:
            # Linux
            candidates = [
                # DroidSansFallback (纯 TrueType，reportlab 兼容最好)
                ("/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf",
                 "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf",
                 None, None),
                ("/usr/share/fonts-droid-fallback/truetype/DroidSansFallback.ttf",
                 "/usr/share/fonts-droid-fallback/truetype/DroidSansFallback.ttf",
                 None, None),
                # Noto Sans CJK (TTC, 需要 subfontIndex)
                ("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
                 "/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc",
                 0, 0),
                ("/usr/share/fonts/opentype/noto/NotoSansCJK-Medium.ttc",
                 "/usr/share/fonts/opentype/noto/NotoSansCJK-Medium.ttc",
                 0, 1),
            ]

        for reg, bld, r_sub, b_sub in candidates:
            if os.path.isfile(reg) and (reg == bld or os.path.isfile(bld)):
                return reg, bld, r_sub, b_sub

        raise FileNotFoundError(
            "未找到可用的中文字体。请安装微软雅黑 (Windows) 或 "
            "DroidSansFallback (Linux) 字体。"
        )

    def export_report_pdf(
        self,
        plan: TestPlan,
        tasks: list[TestTask],
        issues: list[Issue],
        samples: list[Sample],
        filepath: str | None = None,
    ) -> str:
        """导出综合测试报告为 PDF (reportlab)。

        包含：概览统计、任务列表、Issue 列表、样品状态。
        自动检测系统中可用的中文字体（Windows/macOS/Linux）。
        """
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm
        from reportlab.lib.colors import HexColor
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
            PageBreak,
        )
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        from reportlab.pdfgen.canvas import Canvas as _Canvas

        # 跨平台字体检测
        reg_path, bld_path, reg_sub, bld_sub = self._find_cjk_font()
        _FN = "CJK"   # 内部注册名（固定）
        _FN_B = "CJK-Bold"

        # 注册字体（已注册则跳过）
        if _FN not in pdfmetrics.getRegisteredFontNames():
            kw: dict[str, object] = {}
            if reg_sub is not None:
                kw["subfontIndex"] = reg_sub
            pdfmetrics.registerFont(TTFont(_FN, reg_path, **kw))

        if _FN_B not in pdfmetrics.getRegisteredFontNames():
            kw = {}
            if bld_sub is not None:
                kw["subfontIndex"] = bld_sub
            pdfmetrics.registerFont(TTFont(_FN_B, bld_path, **kw))

        pdfmetrics.registerFontFamily(_FN, normal=_FN, bold=_FN_B)

        # 颜色
        _BLUE = HexColor("#2B579A")
        _RED = HexColor("#C0504D")
        _GRAY = HexColor("#646464")
        _LIGHT_GRAY = HexColor("#969696")
        _DARK = HexColor("#323232")

        # 样式
        style_title = ParagraphStyle(
            "Title", fontName=_FN_B, fontSize=24,
            textColor=_BLUE, alignment=TA_CENTER, spaceAfter=8 * mm,
        )
        style_subtitle = ParagraphStyle(
            "Subtitle", fontName=_FN, fontSize=14,
            textColor=_GRAY, alignment=TA_CENTER, spaceAfter=4 * mm,
        )
        style_ts = ParagraphStyle(
            "Timestamp", fontName=_FN, fontSize=10,
            textColor=_LIGHT_GRAY, alignment=TA_CENTER,
        )
        style_section = ParagraphStyle(
            "Section", fontName=_FN_B, fontSize=16,
            textColor=_BLUE, spaceAfter=3 * mm, spaceBefore=6 * mm,
        )
        style_section_red = ParagraphStyle(
            "SectionRed", fontName=_FN_B, fontSize=14,
            textColor=_RED, spaceAfter=2 * mm, spaceBefore=4 * mm,
        )
        style_stat = ParagraphStyle(
            "Stat", fontName=_FN, fontSize=11,
            textColor=_DARK, spaceAfter=1 * mm,
        )
        style_header = ParagraphStyle(
            "Header", fontName=_FN, fontSize=8,
            textColor=_LIGHT_GRAY, alignment=TA_CENTER,
        )

        # 页眉页脚回调
        def _header_footer(canvas: _Canvas, doc: object) -> None:  # noqa: ANN001
            canvas.saveState()
            canvas.setFont(_FN, 8)
            canvas.setFillColor(_LIGHT_GRAY)
            canvas.drawRightString(A4[0] - 20 * mm, A4[1] - 12 * mm,
                                   "ReliaTrack — 可靠性测试报告")
            canvas.drawCentredString(A4[0] / 2, 12 * mm,
                                     f"第 {canvas.getPageNumber()} 页")
            canvas.restoreState()

        out = filepath or str(
            self._ensure_dir() / f"测试报告_{plan.name}_{datetime.now():%Y%m%d_%H%M}.pdf"
        )
        doc = SimpleDocTemplate(
            out, pagesize=A4,
            topMargin=18 * mm, bottomMargin=18 * mm,
            leftMargin=15 * mm, rightMargin=15 * mm,
        )

        story: list[object] = []

        # ── 封面 ──
        story.append(Spacer(1, 40 * mm))
        story.append(Paragraph("可靠性测试报告", style_title))
        story.append(Paragraph(
            f"计划: {plan.name}", style_subtitle))
        story.append(Paragraph(
            f"测试标准: {plan.test_standard or '—'}", style_subtitle))
        story.append(Paragraph(
            f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}", style_ts))

        # ── 概览 ──
        story.append(PageBreak())
        story.append(Paragraph("概览统计", style_section))

        total = len(tasks)
        completed = sum(1 for t in tasks if t.status == "completed")
        in_progress = sum(1 for t in tasks if t.status == "in_progress")
        pending = sum(1 for t in tasks if t.status == "pending")
        total_days = max((t.start_day + t.duration for t in tasks), default=0)
        open_issues = sum(1 for i in issues if i.status in ("open", "analyzing"))
        in_stock = sum(1 for s in samples if s.status == "in_stock")

        stats = [
            f"总任务数: {total}",
            f"已完成: {completed}  |  进行中: {in_progress}  |  待开始: {pending}",
            f"总工期: {total_days} 个工作日",
            f"未关闭 Issue: {open_issues} / {len(issues)}",
            f"在库样品: {in_stock} / {len(samples)}",
        ]
        for s in stats:
            story.append(Paragraph(s, style_stat))

        # ── 任务列表 ──
        story.append(Spacer(1, 6 * mm))
        story.append(Paragraph("测试任务", style_section))

        # 表头
        task_headers = ["#", "名称", "类别", "工期", "开始", "进度", "状态", "优先级"]
        task_col_widths = [18, 130, 55, 40, 40, 40, 50, 40]
        header_row = [Paragraph(h, ParagraphStyle(
            "TH", fontName=_FN_B, fontSize=9,
            textColor=HexColor("#FFFFFF"), alignment=TA_CENTER,
        )) for h in task_headers]

        cell_style = ParagraphStyle(
            "Cell", fontName=_FN, fontSize=8,
            textColor=_DARK, alignment=TA_CENTER,
        )

        task_data = [header_row]
        for idx, task in enumerate(tasks, 1):
            cat = self.CATEGORY_MAP.get(task.category, task.category)
            status = self.STATUS_MAP.get(task.status, task.status)
            task_data.append([
                Paragraph(str(idx), cell_style),
                Paragraph(task.name[:25], cell_style),
                Paragraph(cat, cell_style),
                Paragraph(str(task.duration), cell_style),
                Paragraph(f"D{task.start_day}", cell_style),
                Paragraph(f"{task.progress:.0f}%", cell_style),
                Paragraph(status, cell_style),
                Paragraph(str(task.priority), cell_style),
            ])

        task_table = Table(task_data, colWidths=task_col_widths)
        task_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), _BLUE),
            ("TEXTCOLOR", (0, 0), (-1, 0), HexColor("#FFFFFF")),
            ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#CCCCCC")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1),
             [HexColor("#FFFFFF"), HexColor("#F5F7FA")]),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ]))
        story.append(task_table)

        # ── Issue 列表 ──
        if issues:
            story.append(PageBreak())
            story.append(Paragraph("Issue 追踪", style_section_red))

            issue_headers = ["ID", "标题", "严重度", "状态", "优先级", "失效模式"]
            issue_col_widths = [18, 140, 45, 45, 40, 100]
            issue_header_row = [Paragraph(h, ParagraphStyle(
                "IH", fontName=_FN_B, fontSize=9,
                textColor=HexColor("#FFFFFF"), alignment=TA_CENTER,
            )) for h in issue_headers]

            issue_data = [issue_header_row]
            for issue in issues:
                sev = issue.severity
                status = self.STATUS_MAP.get(issue.status, issue.status)
                issue_data.append([
                    Paragraph(str(issue.id), cell_style),
                    Paragraph(issue.title[:30], cell_style),
                    Paragraph(sev, cell_style),
                    Paragraph(status, cell_style),
                    Paragraph(str(issue.priority), cell_style),
                    Paragraph((issue.failure_mode or "")[:20], cell_style),
                ])

            issue_table = Table(issue_data, colWidths=issue_col_widths)
            issue_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), _RED),
                ("TEXTCOLOR", (0, 0), (-1, 0), HexColor("#FFFFFF")),
                ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#CCCCCC")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1),
                 [HexColor("#FFFFFF"), HexColor("#FFF5F5")]),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]))
            story.append(issue_table)

        # ── 生成 PDF ──
        doc.build(story, onFirstPage=_header_footer, onLaterPages=_header_footer)
        return os.path.abspath(out)

    # ── DVP&R 导出 ──────────────────────────────────────────────

    def export_dvpr_pdf(
        self,
        plan: TestPlan,
        tasks: list[TestTask],
        results: list[TestResult],
        issues: list[Issue],
        samples: list[Sample],
        filepath: str | None = None,
    ) -> str:
        """导出 DVP&R (Design Verification Plan & Report) 格式 PDF。

        DVP&R = 设计验证计划与报告，汽车行业标准格式。
        包含：封面、概览统计、DVP&R 矩阵（任务×样品+判定）、
        Issue/FA/CAPA 汇总、校准状态汇总。
        """
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.lib.units import mm
        from reportlab.lib.colors import HexColor
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
            PageBreak,
        )
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.pdfgen.canvas import Canvas as _Canvas

        reg_path, bld_path, reg_sub, bld_sub = self._find_cjk_font()
        _FN = "CJK"
        _FN_B = "CJK-Bold"
        if _FN not in pdfmetrics.getRegisteredFontNames():
            kw: dict[str, object] = {}
            if reg_sub is not None:
                kw["subfontIndex"] = reg_sub
            pdfmetrics.registerFont(TTFont(_FN, reg_path, **kw))
        if _FN_B not in pdfmetrics.getRegisteredFontNames():
            kw = {}
            if bld_sub is not None:
                kw["subfontIndex"] = bld_sub
            pdfmetrics.registerFont(TTFont(_FN_B, bld_path, **kw))
        pdfmetrics.registerFontFamily(_FN, normal=_FN, bold=_FN_B)

        _BLUE = HexColor("#2B579A")
        _RED = HexColor("#C0504D")
        _GREEN = HexColor("#339933")
        _GRAY = HexColor("#646464")
        _LIGHT_GRAY = HexColor("#969696")
        _DARK = HexColor("#323232")
        _PASS_BG = HexColor("#E8F5E9")
        _FAIL_BG = HexColor("#FFEBEE")

        style_title = ParagraphStyle(
            "Title", fontName=_FN_B, fontSize=24,
            textColor=_BLUE, alignment=TA_CENTER, spaceAfter=8 * mm,
        )
        style_subtitle = ParagraphStyle(
            "Subtitle", fontName=_FN, fontSize=14,
            textColor=_GRAY, alignment=TA_CENTER, spaceAfter=4 * mm,
        )
        style_ts = ParagraphStyle(
            "Timestamp", fontName=_FN, fontSize=10,
            textColor=_LIGHT_GRAY, alignment=TA_CENTER,
        )
        style_section = ParagraphStyle(
            "Section", fontName=_FN_B, fontSize=14,
            textColor=_BLUE, spaceAfter=3 * mm, spaceBefore=6 * mm,
        )
        style_stat = ParagraphStyle(
            "Stat", fontName=_FN, fontSize=11,
            textColor=_DARK, spaceAfter=1 * mm,
        )
        cell_style = ParagraphStyle(
            "Cell", fontName=_FN, fontSize=7,
            textColor=_DARK, alignment=TA_CENTER,
        )
        cell_left = ParagraphStyle(
            "CellL", fontName=_FN, fontSize=7,
            textColor=_DARK, alignment=TA_LEFT,
        )
        th_style = ParagraphStyle(
            "TH", fontName=_FN_B, fontSize=7,
            textColor=HexColor("#FFFFFF"), alignment=TA_CENTER,
        )

        def _header_footer(canvas: _Canvas, doc: object) -> None:
            canvas.saveState()
            canvas.setFont(_FN, 7)
            canvas.setFillColor(_LIGHT_GRAY)
            canvas.drawRightString(
                landscape(A4)[0] - 20 * mm, landscape(A4)[1] - 12 * mm,
                "ReliaTrack — DVP&R Report",
            )
            canvas.drawCentredString(
                landscape(A4)[0] / 2, 10 * mm,
                f"Page {canvas.getPageNumber()}",
            )
            canvas.restoreState()

        out = filepath or str(
            self._ensure_dir() / f"DVP&R_{plan.name}_{datetime.now():%Y%m%d_%H%M}.pdf"
        )
        doc = SimpleDocTemplate(
            out, pagesize=landscape(A4),
            topMargin=15 * mm, bottomMargin=15 * mm,
            leftMargin=12 * mm, rightMargin=12 * mm,
        )

        story: list[object] = []

        # ── 封面 ──
        story.append(Spacer(1, 30 * mm))
        story.append(Paragraph("DVP&R", style_title))
        story.append(Paragraph("Design Verification Plan & Report", style_subtitle))
        story.append(Paragraph(f"计划: {plan.name}", style_subtitle))
        story.append(Paragraph(f"测试标准: {plan.test_standard or '—'}", style_subtitle))
        story.append(Paragraph(
            f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}", style_ts))

        # ── 版本号与日期 ──
        style_version = ParagraphStyle(
            "Version", fontName=_FN, fontSize=11,
            textColor=_DARK, alignment=TA_CENTER, spaceAfter=2 * mm,
        )
        version_line = (
            f"版本: V1.0　　　日期: {datetime.now().strftime('%Y-%m-%d')}"
        )
        story.append(Paragraph(version_line, style_version))

        # ── 概览 ──
        story.append(PageBreak())
        story.append(Paragraph("概览统计", style_section))

        total = len(tasks)
        completed = sum(1 for t in tasks if t.status == "completed")
        total_pass = sum(1 for r in results if r.result == "pass")
        total_fail = sum(1 for r in results if r.result == "fail")
        total_results = len(results)
        pass_rate = f"{total_pass / total_results * 100:.1f}%" if total_results else "—"

        for s in [
            f"总任务数: {total}  |  已完成: {completed}",
            f"测试结果: {total_results}  |  通过: {total_pass}  |  失败: {total_fail}  |  通过率: {pass_rate}",
            f"Issue 数: {len(issues)}",
            f"样品数: {len(samples)}",
        ]:
            story.append(Paragraph(s, style_stat))

        # ── DVP&R 矩阵 ──
        story.append(Spacer(1, 4 * mm))
        story.append(Paragraph("DVP&R 矩阵", style_section))

        # 收集所有涉及样品
        sample_ids = sorted({r.sample_id for r in results if r.sample_id})
        sample_map = {s.id: s.sn for s in samples if s.id is not None}

        # 构建 lookup: (task_id, sample_id) → result
        lookup: dict[tuple[int, int], str] = {}
        for r in results:
            if r.task_id and r.sample_id:
                lookup[(r.task_id, r.sample_id)] = r.result

        # 表头
        dvpr_headers = ["#", "测试项", "判定准则", "样品 SN"]
        for sid in sample_ids:
            dvpr_headers.append(sample_map.get(sid, f"#{sid}"))
        dvpr_headers.append("结论")

        header_row = [Paragraph(h, th_style) for h in dvpr_headers]
        dvpr_data = [header_row]

        for idx, task in enumerate(tasks, 1):
            row = [
                Paragraph(str(idx), cell_style),
                Paragraph((task.name or "")[:20], cell_left),
                Paragraph((task.accept_criteria or "")[:20], cell_left),
            ]
            # 样品 SN 列 — 列出本任务关联的样品编号
            task_sample_ids = sorted({
                r.sample_id for r in results
                if r.task_id == task.id and r.sample_id
            })
            task_sns = ", ".join(
                sample_map.get(sid, f"#{sid}") for sid in task_sample_ids
            )
            row.append(Paragraph((task_sns or "—")[:25], cell_left))
            task_pass = 0
            task_fail = 0
            for sid in sample_ids:
                res = lookup.get((task.id, sid), "")
                if res == "pass":
                    row.append(Paragraph("P", cell_style))
                    task_pass += 1
                elif res == "fail":
                    row.append(Paragraph("F", cell_style))
                    task_fail += 1
                elif res == "conditional":
                    row.append(Paragraph("C", cell_style))
                elif res == "skip":
                    row.append(Paragraph("S", cell_style))
                else:
                    row.append(Paragraph("—", cell_style))
            # 结论
            if task_fail > 0:
                row.append(Paragraph("FAIL", cell_style))
            elif task_pass > 0 and task_fail == 0:
                row.append(Paragraph("PASS", cell_style))
            else:
                row.append(Paragraph("—", cell_style))
            dvpr_data.append(row)

        # 列宽
        n_samples = len(sample_ids)
        page_w = landscape(A4)[0] - 24 * mm
        fixed_cols = 18 + 80 + 65 + 60 + 35  # # + name + criteria + SN + conclusion
        sample_col_w = max(35, (page_w - fixed_cols) / max(n_samples, 1))
        dvpr_widths = [18, 80, 65, 60] + [sample_col_w] * n_samples + [35]

        dvpr_table = Table(dvpr_data, colWidths=dvpr_widths)
        table_styles = [
            ("BACKGROUND", (0, 0), (-1, 0), _BLUE),
            ("TEXTCOLOR", (0, 0), (-1, 0), HexColor("#FFFFFF")),
            ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#CCCCCC")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1),
             [HexColor("#FFFFFF"), HexColor("#F5F7FA")]),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING", (0, 0), (-1, -1), 2),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
        ]
        # 着色 pass/fail 单元格
        for row_idx, task in enumerate(tasks, 1):
            for col_idx, sid in enumerate(sample_ids):
                res = lookup.get((task.id, sid), "")
                if res == "fail":
                    table_styles.append(("BACKGROUND", (col_idx + 4, row_idx), (col_idx + 4, row_idx), _FAIL_BG))
                elif res == "pass":
                    table_styles.append(("BACKGROUND", (col_idx + 4, row_idx), (col_idx + 4, row_idx), _PASS_BG))

        dvpr_table.setStyle(TableStyle(table_styles))
        story.append(dvpr_table)

        # ── Issue 汇总 ──
        if issues:
            story.append(PageBreak())
            story.append(Paragraph("Issue 追踪汇总", style_section))
            issue_headers = ["ID", "标题", "严重度", "状态", "失效模式"]
            issue_data = [[Paragraph(h, th_style) for h in issue_headers]]
            for issue in issues:
                issue_data.append([
                    Paragraph(str(issue.id), cell_style),
                    Paragraph((issue.title or "")[:30], cell_left),
                    Paragraph(issue.severity, cell_style),
                    Paragraph(self.STATUS_MAP.get(issue.status, issue.status), cell_style),
                    Paragraph((issue.failure_mode or "")[:20], cell_left),
                ])
            issue_table = Table(issue_data, colWidths=[25, 160, 50, 50, 120])
            issue_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), _RED),
                ("TEXTCOLOR", (0, 0), (-1, 0), HexColor("#FFFFFF")),
                ("GRID", (0, 0), (-1, -1), 0.5, HexColor("#CCCCCC")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING", (0, 0), (-1, -1), 2),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
            ]))
            story.append(issue_table)

        # ── 签字栏 ──
        story.append(Spacer(1, 12 * mm))
        sig_label_style = ParagraphStyle(
            "SigLabel", fontName=_FN_B, fontSize=9,
            textColor=_DARK, alignment=TA_CENTER,
        )
        sig_line_style = ParagraphStyle(
            "SigLine", fontName=_FN, fontSize=9,
            textColor=_LIGHT_GRAY, alignment=TA_CENTER,
        )
        sig_roles = ["编制", "审核", "批准"]
        sig_cells = []
        for role in sig_roles:
            sig_cells.append([
                Paragraph(role, sig_label_style),
                Spacer(1, 10 * mm),
                Paragraph("________________________", sig_line_style),
            ])
        sig_table = Table(
            [sig_cells],
            colWidths=[page_w / 3] * 3,
        )
        sig_table.setStyle(TableStyle([
            ("BOX", (0, 0), (0, 0), 0.5, _GRAY),
            ("BOX", (1, 0), (1, 0), 0.5, _GRAY),
            ("BOX", (2, 0), (2, 0), 0.5, _GRAY),
            ("VALIGN", (0, 0), (-1, -1), "BOTTOM"),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        story.append(sig_table)

        doc.build(story, onFirstPage=_header_footer, onLaterPages=_header_footer)
        return os.path.abspath(out)

    # ── 8D 报告导出 ──────────────────────────────────────────────

    def export_8d_pdf(
        self,
        issue: Issue,
        fa_records: list[FARecord] | None = None,
        capa_records: list[CAPARecord] | None = None,
        filepath: str | None = None,
    ) -> str:
        """导出 8D Problem Solving Report 为 PDF。

        包含：基本信息表、D1-D8 八个章节、底部签字栏。
        用 Issue 已有数据填充可匹配的部分，其余留空供用户手写。
        """
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm
        from reportlab.lib.colors import HexColor
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        )
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.pdfgen.canvas import Canvas as _Canvas

        reg_path, bld_path, reg_sub, bld_sub = self._find_cjk_font()
        _FN = "CJK"
        _FN_B = "CJK-Bold"
        if _FN not in pdfmetrics.getRegisteredFontNames():
            kw: dict[str, object] = {}
            if reg_sub is not None:
                kw["subfontIndex"] = reg_sub
            pdfmetrics.registerFont(TTFont(_FN, reg_path, **kw))
        if _FN_B not in pdfmetrics.getRegisteredFontNames():
            kw2: dict[str, object] = {}
            if bld_sub is not None:
                kw2["subfontIndex"] = bld_sub
            pdfmetrics.registerFont(TTFont(_FN_B, bld_path, **kw2))
        pdfmetrics.registerFontFamily(_FN, normal=_FN, bold=_FN_B)

        _BLUE = HexColor("#2B579A")
        _RED = HexColor("#C0504D")
        _GRAY = HexColor("#646464")
        _LIGHT_GRAY = HexColor("#969696")
        _DARK = HexColor("#323232")
        _SECTION_BG = HexColor("#E8EDF5")

        style_title = ParagraphStyle(
            "Title8D", fontName=_FN_B, fontSize=22,
            textColor=_BLUE, alignment=TA_CENTER, spaceAfter=6 * mm,
        )
        style_subtitle = ParagraphStyle(
            "Sub8D", fontName=_FN, fontSize=12,
            textColor=_GRAY, alignment=TA_CENTER, spaceAfter=4 * mm,
        )
        style_section_label = ParagraphStyle(
            "SecL", fontName=_FN_B, fontSize=11,
            textColor=_BLUE, spaceAfter=1 * mm, spaceBefore=2 * mm,
        )
        style_body = ParagraphStyle(
            "Body8D", fontName=_FN, fontSize=9,
            textColor=_DARK, alignment=TA_LEFT, leading=14,
        )
        style_cell = ParagraphStyle(
            "Cell8D", fontName=_FN, fontSize=9,
            textColor=_DARK, alignment=TA_LEFT, leading=13,
        )
        style_cell_center = ParagraphStyle(
            "CellC8D", fontName=_FN, fontSize=9,
            textColor=_DARK, alignment=TA_CENTER, leading=13,
        )
        style_cell_bold = ParagraphStyle(
            "CellB8D", fontName=_FN_B, fontSize=9,
            textColor=_BLUE, alignment=TA_CENTER, leading=13,
        )
        style_blank = ParagraphStyle(
            "Blank8D", fontName=_FN, fontSize=9,
            textColor=_LIGHT_GRAY, alignment=TA_LEFT, leading=13,
        )
        style_sig_label = ParagraphStyle(
            "SigL8D", fontName=_FN_B, fontSize=9,
            textColor=_DARK, alignment=TA_CENTER,
        )
        style_sig_line = ParagraphStyle(
            "SigLine8D", fontName=_FN, fontSize=9,
            textColor=_LIGHT_GRAY, alignment=TA_CENTER,
        )

        def _header_footer(canvas: _Canvas, doc: object) -> None:
            canvas.saveState()
            canvas.setFont(_FN, 8)
            canvas.setFillColor(_LIGHT_GRAY)
            canvas.drawRightString(
                A4[0] - 20 * mm, A4[1] - 12 * mm,
                "ReliaTrack — 8D Problem Solving Report",
            )
            canvas.drawCentredString(
                A4[0] / 2, 12 * mm,
                f"Page {canvas.getPageNumber()}",
            )
            canvas.restoreState()

        out = filepath or str(
            self._ensure_dir() / f"8D_Report_Issue{issue.id}_{datetime.now():%Y%m%d_%H%M}.pdf"
        )
        doc = SimpleDocTemplate(
            out, pagesize=A4,
            topMargin=18 * mm, bottomMargin=18 * mm,
            leftMargin=15 * mm, rightMargin=15 * mm,
        )

        page_w = A4[0] - 30 * mm  # usable width

        story: list[object] = []

        # ── 标题 ──
        story.append(Spacer(1, 8 * mm))
        story.append(Paragraph("8D Problem Solving Report", style_title))
        story.append(Paragraph(
            f"Issue #{issue.id} — {issue.title}", style_subtitle))
        story.append(Paragraph(
            f"报告日期: {datetime.now().strftime('%Y-%m-%d')}", style_subtitle))

        # ── 基本信息表 ──
        severity_labels = {
            "critical": "Critical (致命)",
            "major": "Major (严重)",
            "minor": "Minor (一般)",
            "cosmetic": "Cosmetic (外观)",
        }
        sev_text = severity_labels.get(issue.severity, issue.severity)
        status_text = self.STATUS_MAP.get(issue.status, issue.status)

        info_data = [
            [
                Paragraph("Issue 编号", style_cell_bold),
                Paragraph(str(issue.id), style_cell_center),
                Paragraph("严重度", style_cell_bold),
                Paragraph(sev_text, style_cell_center),
            ],
            [
                Paragraph("标题", style_cell_bold),
                Paragraph(issue.title, style_cell),
                Paragraph("状态", style_cell_bold),
                Paragraph(status_text, style_cell_center),
            ],
            [
                Paragraph("报告日期", style_cell_bold),
                Paragraph(datetime.now().strftime("%Y-%m-%d"), style_cell_center),
                Paragraph("优先级", style_cell_bold),
                Paragraph(str(issue.priority), style_cell_center),
            ],
        ]
        info_table = Table(info_data, colWidths=[page_w * 0.15, page_w * 0.35, page_w * 0.15, page_w * 0.35])
        info_table.setStyle(TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.5, _GRAY),
            ("BACKGROUND", (0, 0), (0, -1), _SECTION_BG),
            ("BACKGROUND", (2, 0), (2, -1), _SECTION_BG),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ]))
        story.append(info_table)

        # ── D1-D8 章节 ──
        d_sections: list[tuple[str, str, str]] = [
            ("D1", "团队组建 (Establish the Team)", "(手写区)"),
            ("D2", "问题描述 (Describe the Problem)", issue.description or ""),
            ("D3", "临时遏制措施 (Interim Containment Actions)", "(手写区)"),
            ("D4", "根因分析 (Root Cause Analysis)", self._build_d4_content(issue, fa_records)),
            ("D5", "纠正措施 (Corrective Actions)", issue.resolution or ""),
            ("D6", "实施验证 (Implement & Validate)", self._build_d6_content(capa_records)),
            ("D7", "预防再发 (Prevent Recurrence)", "(手写区)"),
            ("D8", "结论与签字 (Congratulate the Team)", "(签字区)"),
        ]

        for d_label, d_title, d_content in d_sections:
            story.append(Spacer(1, 3 * mm))

            # D 章节标题行
            header_data = [[
                Paragraph(f"<b>{d_label}</b>", ParagraphStyle(
                    "DH", fontName=_FN_B, fontSize=10,
                    textColor=HexColor("#FFFFFF"), alignment=TA_CENTER,
                )),
                Paragraph(f"<b>{d_title}</b>", ParagraphStyle(
                    "DT", fontName=_FN_B, fontSize=10,
                    textColor=HexColor("#FFFFFF"), alignment=TA_LEFT,
                )),
            ]]
            header_table = Table(header_data, colWidths=[page_w * 0.12, page_w * 0.88])
            header_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), _BLUE),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ]))
            story.append(header_table)

            # D 内容行
            content_text = d_content if d_content else ""
            content_style = style_blank if content_text.startswith("(") else style_cell
            content_data = [[
                Paragraph("", style_cell),
                Paragraph(content_text or "", content_style),
            ]]
            content_table = Table(content_data, colWidths=[page_w * 0.12, page_w * 0.88])
            min_h = 50 if content_text.startswith("(") else max(30, len(content_text) // 2)
            content_table.setStyle(TableStyle([
                ("BOX", (0, 0), (-1, -1), 0.5, _GRAY),
                ("LINEAFTER", (0, 0), (0, -1), 0.5, _GRAY),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("ROWHEIGHTS", (0, 0), (-1, -1), min_h),
            ]))
            story.append(content_table)

        # ── 底部签字栏 ──
        story.append(Spacer(1, 8 * mm))
        sig_roles = ["编制 (Prepared)", "审核 (Reviewed)", "批准 (Approved)"]
        sig_cells = []
        for role in sig_roles:
            sig_cells.append([
                Paragraph(f"<b>{role}</b>", style_sig_label),
                Spacer(1, 10 * mm),
                Paragraph("________________________", style_sig_line),
            ])
        sig_table = Table(
            [sig_cells],
            colWidths=[page_w / 3] * 3,
        )
        sig_table.setStyle(TableStyle([
            ("BOX", (0, 0), (0, 0), 0.5, _GRAY),
            ("BOX", (1, 0), (1, 0), 0.5, _GRAY),
            ("BOX", (2, 0), (2, 0), 0.5, _GRAY),
            ("VALIGN", (0, 0), (-1, -1), "BOTTOM"),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        story.append(sig_table)

        doc.build(story, onFirstPage=_header_footer, onLaterPages=_header_footer)
        return os.path.abspath(out)

    @staticmethod
    def _build_d4_content(issue: Issue, fa_records: list[FARecord] | None) -> str:
        """构建 D4 根因分析内容。"""
        parts: list[str] = []
        if issue.root_cause:
            parts.append(f"根因: {issue.root_cause}")
        if issue.failure_mode:
            parts.append(f"失效模式: {issue.failure_mode}")
        if fa_records:
            parts.append("\nFA 分析记录摘要:")
            for rec in fa_records:
                confirmed_labels = {0: "待定", 1: "确认", 2: "排除"}
                status = confirmed_labels.get(rec.confirmed, "待定")
                parts.append(
                    f"  Step {rec.step_no}: {rec.step_title or ''}"
                    f" — {rec.method or ''}"
                    f" — 发现: {rec.findings or '—'}"
                    f" — 可能原因: {rec.possible_cause or '—'}"
                    f" [{status}]"
                )
        return "\n".join(parts) if parts else ""

    @staticmethod
    def _build_d6_content(capa_records: list[CAPARecord] | None) -> str:
        """构建 D6 实施验证内容。"""
        if not capa_records:
            return ""
        parts: list[str] = []
        for rec in capa_records:
            status_labels = {
                "pending": "待执行", "in_progress": "进行中",
                "completed": "已完成", "verified": "已验证",
            }
            status = status_labels.get(rec.status, rec.status)
            parts.append(f"措施: {rec.action}")
            parts.append(f"  状态: {status}")
            if rec.verification_result:
                parts.append(f"  验证结果: {rec.verification_result}")
        return "\n".join(parts)

    # ── Word 导出 ──────────────────────────────────────────────

    def export_to_word(
        self,
        plan: TestPlan,
        tasks: list[TestTask],
        issues: list[Issue],
        samples: list[Sample],
        filepath: str | None = None,
    ) -> str:
        """导出综合测试报告为 Word (.docx)。

        包含：封面标题、项目信息表格、任务列表表格、Issue 列表表格、样品列表表格。
        """
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn, nsdecls
        from docx.oxml import parse_xml

        _f = self.get_cjk_font()

        doc = Document()

        # ── 样式设置（全局一次，子元素自动继承） ──
        style = doc.styles["Normal"]
        font = style.font
        font.name = _f
        font.size = Pt(10)
        font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        style.element.rPr.rFonts.set(qn("w:eastAsia"), _f)
        pf = style.paragraph_format
        pf.space_after = Pt(4)
        pf.space_before = Pt(2)

        # ── 快速写入单元格（直接操作 lxml 避免 parse_xml 开销） ──
        from lxml import etree
        _ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        _NSMAP = {"w": _ns}

        def _make(tag, **attrib):
            return etree.SubElement(etree.Element("dummy"), f"{{{_ns}}}{tag}").makeelement(
                f"{{{_ns}}}{tag}", {f"{{{_ns}}}{k}": v for k, v in attrib.items()}
            )

        def _fill_cell(ct_tc, text, bold=False, size=9, color=None,
                        shade=None, align=None):
            """直接操作 CT_Tc XML 元素，避免 _Cell 包装开销。"""
            # 清除旧段落
            for p in ct_tc.findall(qn("w:p")):
                ct_tc.remove(p)
            # 构建 <w:p><w:r><w:rPr>...<w:t>xml:space="preserve">text</w:t></w:r></w:p>
            p = etree.SubElement(ct_tc, f"{{{_ns}}}p")
            if align == "center":
                pPr = etree.SubElement(p, f"{{{_ns}}}pPr")
                etree.SubElement(pPr, f"{{{_ns}}}jc", attrib={f"{{{_ns}}}val": "center"})
            r = etree.SubElement(p, f"{{{_ns}}}r")
            rPr = etree.SubElement(r, f"{{{_ns}}}rPr")
            rFonts = etree.SubElement(rPr, f"{{{_ns}}}rFonts")
            rFonts.set(f"{{{_ns}}}ascii", _f)
            rFonts.set(f"{{{_ns}}}eastAsia", _f)
            rFonts.set(f"{{{_ns}}}hAnsi", _f)
            rPr.append(_make("sz", val=str(size * 2)))
            if bold:
                rPr.append(_make("b"))
            if color:
                rPr.append(_make("color", val=color))
            t = etree.SubElement(r, f"{{{_ns}}}t")
            t.text = text
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            if shade:
                tcPr = ct_tc.get_or_add_tcPr()
                tcPr.append(_make("shd", fill=shade, val="clear"))

        def _fill_row_cells(tr, values, bold=False, color=None,
                            shade=None, center_cols=frozenset()):
            """批量填充一行，跳过 _Cell 构造。"""
            tc_elems = tr.findall(qn("w:tc"))
            for j, val in enumerate(values):
                tc = tc_elems[j] if j < len(tc_elems) else None
                if tc is not None:
                    _fill_cell(tc, str(val), bold=bold, color=color,
                               shade=shade if j == 0 or shade else None,
                               align="center" if j in center_cols else None)

        # ── 标题 ──
        title = doc.add_heading(level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("可靠性测试报告")
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)

        # 副标题
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub.add_run(f"计划: {plan.name}  |  测试标准: {plan.test_standard or '—'}")
        sub_run.font.size = Pt(12)
        sub_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        # 导出时间
        ts = doc.add_paragraph()
        ts.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ts_run = ts.add_run(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        ts_run.font.size = Pt(10)
        ts_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # ── 概览统计 ──
        doc.add_heading("概览统计", level=2)
        total = len(tasks)
        completed = sum(1 for t in tasks if t.status == "completed")
        in_progress = sum(1 for t in tasks if t.status == "in_progress")
        pending = sum(1 for t in tasks if t.status == "pending")
        total_days = max((t.start_day + t.duration for t in tasks), default=0)
        open_issues = sum(1 for i in issues if i.status in ("open", "analyzing"))
        in_stock = sum(1 for s in samples if s.status == "in_stock")

        stats_lines = [
            f"总任务数: {total}    已完成: {completed}    进行中: {in_progress}    待开始: {pending}",
            f"总工期: {total_days} 个工作日",
            f"未关闭 Issue: {open_issues} / {len(issues)}",
            f"在库样品: {in_stock} / {len(samples)}",
        ]
        for line in stats_lines:
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(2)

        # ── 项目信息表格 ──
        doc.add_heading("项目信息", level=2)
        info_table = doc.add_table(rows=5, cols=2, style="Table Grid")
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        info_data = [
            ("计划名称", plan.name),
            ("测试标准", plan.test_standard or "—"),
            ("开始日期", plan.start_date or "—"),
            ("结束日期", plan.end_date or "—"),
            ("计划状态", self.STATUS_MAP.get(plan.status, plan.status)),
        ]
        for i, (label, value) in enumerate(info_data):
            tr = info_table.rows[i]._tr
            tc_elems = tr.findall(qn("w:tc"))
            _fill_cell(tc_elems[0], label, bold=True, size=10, shade="E8EDF5")
            _fill_cell(tc_elems[1], value, size=10)

        # ── 任务列表表格 ──
        doc.add_heading("测试任务", level=2)
        task_headers = ["#", "名称", "类别", "状态", "天数", "设备", "技术员", "进度"]
        task_table = doc.add_table(rows=1 + len(tasks), cols=len(task_headers), style="Table Grid")
        task_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _fill_row_cells(task_table.rows[0]._tr, task_headers, bold=True,
                        color="FFFFFF", shade="2B579A", center_cols=set(range(len(task_headers))))

        for idx, task in enumerate(tasks, 1):
            equipment_name = f"ID:{task.equipment_id}" if task.equipment_id else "—"
            technician_name = f"ID:{task.technician_id}" if task.technician_id else "—"
            _fill_row_cells(task_table.rows[idx]._tr, [
                idx, task.name,
                self.CATEGORY_MAP.get(task.category, task.category),
                self.STATUS_MAP.get(task.status, task.status),
                task.duration, equipment_name, technician_name,
                f"{task.progress:.0f}%",
            ], center_cols={0})

        # ── Issue 列表表格 ──
        if issues:
            doc.add_heading("Issue 追踪", level=2)
            issue_headers = ["#", "标题", "优先级", "状态", "根因"]
            issue_table = doc.add_table(
                rows=1 + len(issues), cols=len(issue_headers), style="Table Grid"
            )
            issue_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            _fill_row_cells(issue_table.rows[0]._tr, issue_headers, bold=True,
                            color="FFFFFF", shade="C0504D", center_cols=set(range(len(issue_headers))))

            for idx, issue in enumerate(issues, 1):
                _fill_row_cells(issue_table.rows[idx]._tr, [
                    idx, issue.title, issue.priority,
                    self.STATUS_MAP.get(issue.status, issue.status),
                    (issue.root_cause or "")[:80] if issue.root_cause else "",
                ], center_cols={0})

        # ── 样品列表表格 ──
        if samples:
            doc.add_heading("样品台账", level=2)
            sample_headers = ["#", "SN", "批次号", "规格型号", "状态"]
            sample_table = doc.add_table(
                rows=1 + len(samples), cols=len(sample_headers), style="Table Grid"
            )
            sample_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            _fill_row_cells(sample_table.rows[0]._tr, sample_headers, bold=True,
                            color="FFFFFF", shade="4F81BD", center_cols=set(range(len(sample_headers))))

            for idx, s in enumerate(samples, 1):
                _fill_row_cells(sample_table.rows[idx]._tr, [
                    idx, s.sn, s.batch_no, s.spec or "",
                    self.STATUS_MAP.get(s.status, s.status),
                ], center_cols={0})

        # ── 保存 ──
        out = filepath or str(
            self._ensure_dir() / f"测试报告_{plan.name}_{datetime.now():%Y%m%d_%H%M}.docx"
        )
        doc.save(out)
        return os.path.abspath(out)

    # ── 辅助方法 ──

    @staticmethod
    def _set_cell_shading(cell, hex_color: str) -> None:
        """设置 Word 单元格背景色。"""
        from docx.oxml.ns import qn, nsdecls
        from docx.oxml import parse_xml

        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'
        )
        cell._tc.get_or_add_tcPr().append(shading)
