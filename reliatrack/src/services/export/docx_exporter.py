"""导出服务 — Word 导出器。

包含：综合报告 Word、DVP&R Word、8D Report Word（python-docx）。
"""

from __future__ import annotations

import logging
import os
from datetime import datetime
from pathlib import Path
from typing import TYPE_CHECKING

from src.services.export.export_utils import (
    CATEGORY_MAP, STATUS_MAP, _validate_output_path, logger,
    excel_styles, excel_save, excel_write_headers, excel_write_row,
    _judge_conclusion,
)
from src.constants import RESOLUTION_LABELS

if TYPE_CHECKING:
    from src.models.issue import Issue, FARecord, CAPARecord
    from src.models.sample import Sample
    from src.models.test_plan import TestPlan, TestTask, TestResult

import platform
import subprocess

_CJK_FONT_LOCAL = None


def _get_cjk_font():
    global _CJK_FONT_LOCAL
    if _CJK_FONT_LOCAL is not None:
        return _CJK_FONT_LOCAL
    system = platform.system()
    if system == "Windows":
        _CJK_FONT_LOCAL = "Microsoft YaHei"
    elif system == "Darwin":
        _CJK_FONT_LOCAL = "PingFang SC"
    else:
        for candidate in ("Noto Sans CJK SC", "WenQuanYi Micro Hei", "Droid Sans Fallback"):
            try:
                r = subprocess.run(["fc-match", "-f", "%{family}", candidate],
                                   capture_output=True, text=True, timeout=3)
                if r.returncode == 0 and r.stdout.strip():
                    _CJK_FONT_LOCAL = r.stdout.strip()
                    break
            except Exception:
                continue
        if _CJK_FONT_LOCAL is None:
            _CJK_FONT_LOCAL = "Noto Sans CJK SC"
    return _CJK_FONT_LOCAL


def _set_table_border(table):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "999999")
        borders.append(el)
    tblPr.append(borders)


def _fill_cell_local(ct_tc, text, bold=False, size=9, color=None, shade=None, align=None):
    for p in ct_tc.findall(qn("w:p")):
        ct_tc.remove(p)
    p = etree_SubElement(ct_tc, "w:p")
    if align == "center":
        pPr = etree_SubElement(p, "w:pPr")
        etree_SubElement(pPr, "w:jc").set(qn("w:val"), "center")
    r = etree_SubElement(p, "w:r")
    rPr = etree_SubElement(r, "w:rPr")
    rFonts = etree_SubElement(rPr, "w:rFonts")
    _f = _get_cjk_font()
    rFonts.set(qn("w:ascii"), _f)
    rFonts.set(qn("w:eastAsia"), _f)
    rFonts.set(qn("w:hAnsi"), _f)
    if bold:
        etree_SubElement(rPr, "w:b")
    if color:
        etree_SubElement(rPr, "w:color").set(qn("w:val"), color)
    etree_SubElement(rPr, "w:sz").set(qn("w:val"), str(size * 2))
    from docx.oxml.ns import nsdecls
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    t = etree_SubElement(r, "w:t")
    t.text = str(text) if text else ""
    t.set(f"{{{ns}}}space", "preserve")
    if shade:
        tcPr = ct_tc.get_or_add_tcPr()
        shd = etree_SubElement(tcPr, "w:shd")
        shd.set(qn("w:fill"), shade)
        shd.set(qn("w:val"), "clear")


def _fill_row_cells_local(tr, values, bold=False, color=None, shade=None, center_cols=frozenset()):
    from docx.oxml.ns import qn
    tc_elems = tr.findall(qn("w:tc"))
    for j, val in enumerate(values):
        tc = tc_elems[j] if j < len(tc_elems) else None
        if tc is not None:
            _fill_cell_local(
                tc, str(val), bold=bold, color=color,
                shade=shade if j == 0 else shade,
                align="center" if j in center_cols else None,
            )


def etree_SubElement(parent, tag):
    from lxml import etree as _etree
    return _etree.SubElement(parent, f"{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}{tag}")


def export_to_word(
    output_dir: Path,
    plan: TestPlan,
    tasks: list[TestTask],
    issues: list[Issue],
    samples: list[Sample],
    filepath: str | None = None,
    results: list | None = None,
) -> str:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import parse_xml

    _f = _get_cjk_font()
    _BLUE = RGBColor(0x2B, 0x57, 0x9A)

    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = _f
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), _f)
    pf = style.paragraph_format
    pf.space_after = Pt(4)
    pf.space_before = Pt(2)

    from lxml import etree
    _ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

    def _make_ns(tag, **attrib):
        return etree.SubElement(etree.Element("dummy"), f"{{{_ns}}}{tag}").makeelement(
            f"{{{_ns}}}{tag}", {f"{{{_ns}}}{k}": v for k, v in attrib.items()})

    def _fill_cell(ct_tc, text, bold=False, size=9, color=None, shade=None, align=None):
        for p in ct_tc.findall(qn("w:p")):
            ct_tc.remove(p)
        p = etree.SubElement(ct_tc, f"{{{_ns}}}p")
        if align == "center":
            pPr = etree.SubElement(p, f"{{{_ns}}}pPr")
            etree.SubElement(pPr, f"{{{_ns}}}jc", attrib={f"{{{_ns}}}val": "center"})
        r = etree.SubElement(p, f"{{{_ns}}}r")
        rPr = etree.SubElement(r, f"{{{_ns}}}rPr")
        rFonts = etree.SubElement(rPr, f"{{{_ns}}}rFonts")
        rFonts.set(f"{{{_ns}}}ascii", _f)
        rFonts.set(f"{{{_ns}}}eastAsia", _f)
        rFonts.set(f"{{{_ns}}}hAnsi", _f)
        rPr.append(_make_ns("sz", val=str(size * 2)))
        if bold:
            rPr.append(_make_ns("b"))
        if color:
            rPr.append(_make_ns("color", val=color))
        t = etree.SubElement(r, f"{{{_ns}}}t")
        t.text = str(text) if text else ""
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        if shade:
            tcPr = ct_tc.get_or_add_tcPr()
            shd = etree.SubElement(tcPr, f"{{{_ns}}}shd")
            shd.set(f"{{{_ns}}}fill", shade)
            shd.set(f"{{{_ns}}}val", "clear")

    def _fill_row_cells(tr, values, bold=False, color=None, shade=None, center_cols=frozenset()):
        tc_elems = tr.findall(qn("w:tc"))
        for j, val in enumerate(values):
            tc = tc_elems[j] if j < len(tc_elems) else None
            if tc is not None:
                _fill_cell(tc, str(val), bold=bold, color=color,
                           shade=shade if j == 0 or shade else None,
                           align="center" if j in center_cols else None)

    # ── 标题 ──
    title = doc.add_heading(level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("可靠性测试报告")
    run.font.size = Pt(24)
    run.font.color.rgb = _BLUE

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(f"计划: {plan.name}  |  测试标准: {plan.test_standard or '—'}")
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    ts = doc.add_paragraph()
    ts.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ts_run = ts.add_run(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    ts_run.font.size = Pt(10)
    ts_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # ── 概览统计 ──
    doc.add_heading("概览统计", level=2)
    total = len(tasks)
    completed = sum(1 for t in tasks if t.status == "completed")
    in_progress = sum(1 for t in tasks if t.status == "in_progress")
    pending = sum(1 for t in tasks if t.status == "pending")
    total_days = max((t.start_day + t.duration for t in tasks), default=0)
    open_issues = sum(1 for i in issues if i.status in ("open", "analyzing"))
    in_stock = sum(1 for s in samples if s.status == "in_stock")

    for line in [
        f"总任务数: {total}    已完成: {completed}    进行中: {in_progress}    待开始: {pending}",
        f"总工期: {total_days} 个工作日",
        f"未关闭 Issue: {open_issues} / {len(issues)}",
        f"在库样品: {in_stock} / {len(samples)}",
    ]:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(2)

    # ── 项目信息表格 ──
    doc.add_heading("项目信息", level=2)
    info_table = doc.add_table(rows=5, cols=2, style="Table Grid")
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate([
        ("计划名称", plan.name),
        ("测试标准", plan.test_standard or "—"),
        ("开始日期", plan.start_date or "—"),
        ("结束日期", plan.end_date or "—"),
        ("计划状态", STATUS_MAP.get(plan.status, plan.status)),
    ]):
        tr = info_table.rows[i]._tr
        tc_elems = tr.findall(qn("w:tc"))
        _fill_cell(tc_elems[0], label, bold=True, size=10, shade="E8EDF5")
        _fill_cell(tc_elems[1], value, size=10)

    # ── 任务列表表格 ──
    doc.add_heading("测试任务", level=2)
    task_headers = ["#", "名称", "类别", "状态", "天数", "设备", "技术员", "进度"]
    task_table = doc.add_table(rows=1 + len(tasks), cols=len(task_headers), style="Table Grid")
    task_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _fill_row_cells(task_table.rows[0]._tr, task_headers, bold=True,
                    color="FFFFFF", shade="2B579A", center_cols=set(range(len(task_headers))))

    for idx, task in enumerate(tasks, 1):
        equipment_name = f"ID:{task.equipment_id}" if task.equipment_id else "—"
        technician_name = f"ID:{task.technician_id}" if task.technician_id else "—"
        _fill_row_cells(task_table.rows[idx]._tr, [
            idx, task.name,
            CATEGORY_MAP.get(task.category, task.category),
            STATUS_MAP.get(task.status, task.status),
            task.duration, equipment_name, technician_name,
            f"{task.progress:.0f}%",
        ], center_cols={0})

    # ── Issue 列表表格 ──
    if issues:
        doc.add_heading("Issue 追踪", level=2)
        issue_headers = ["#", "Issue描述", "优先级", "状态", "DRI", "报告人", "解决结果", "根因"]
        issue_table = doc.add_table(
            rows=1 + len(issues), cols=len(issue_headers), style="Table Grid"
        )
        issue_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _fill_row_cells(issue_table.rows[0]._tr, issue_headers, bold=True,
                        color="FFFFFF", shade="C0504D", center_cols=set(range(len(issue_headers))))

        for idx, issue in enumerate(issues, 1):
            resolution_text = RESOLUTION_LABELS.get(issue.resolution, issue.resolution) or ""
            _fill_row_cells(issue_table.rows[idx]._tr, [
                idx, issue.title, issue.priority,
                STATUS_MAP.get(issue.status, issue.status),
                issue.dri_name or "",
                issue.reporter_name or "",
                resolution_text,
                (issue.root_cause or "")[:80] if issue.root_cause else "",
            ], center_cols={0})

    # ── 测试结果汇总 ──
    _results = results or []
    if _results:
        doc.add_heading("测试结果汇总", level=2)
        sample_map = {s.id: s.sn for s in samples if s.id is not None}
        task_map = {t.id: t for t in tasks if t.id is not None}
        total_pass = sum(1 for r in _results if r.result == "pass")
        total_fail = sum(1 for r in _results if r.result == "fail")
        total_conditional = sum(1 for r in _results if r.result == "conditional")
        total_results = len(_results)
        pass_rate = f"{total_pass / total_results * 100:.1f}%" if total_results else "—"

        overall_conclusion = _judge_conclusion(
            total_pass, total_fail, total_conditional, total_results,
            accept_criteria="",
        )
        stat_lines = [
            f"测试结果总数: {total_results}  |  通过: {total_pass}  |  失败: {total_fail}  |  条件通过: {total_conditional}  |  通过率: {pass_rate}",
            f"总体判定结论: {overall_conclusion}",
        ]
        for line in stat_lines:
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(2)

        doc.add_heading("结果明细", level=3)
        res_headers = ["#", "任务名", "样品SN", "结果", "判定"]
        res_table = doc.add_table(
            rows=1 + len(_results), cols=len(res_headers), style="Table Grid"
        )
        res_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _fill_row_cells(res_table.rows[0]._tr, res_headers, bold=True,
                        color="FFFFFF", shade="339933",
                        center_cols=set(range(len(res_headers))))

        for idx, r in enumerate(_results, 1):
            task_name = (task_map[r.task_id].name or "")[:25] if r.task_id and r.task_id in task_map else ""
            sample_sn = sample_map.get(r.sample_id, f"#{r.sample_id}") if r.sample_id else "—"
            result_text = r.result.upper() if r.result else "—"
            task_obj = task_map.get(r.task_id)
            task_pass = sum(1 for x in _results if x.task_id == r.task_id and x.result == "pass")
            task_fail = sum(1 for x in _results if x.task_id == r.task_id and x.result == "fail")
            task_cond = sum(1 for x in _results if x.task_id == r.task_id and x.result == "conditional")
            accept_criteria = task_obj.accept_criteria if task_obj else ""
            conclusion = _judge_conclusion(
                task_pass, task_fail, task_cond, 0,
                accept_criteria=accept_criteria or "",
            )
            if idx > 1 and _results[idx - 2].task_id == r.task_id:
                conclusion = ""
            _fill_row_cells(res_table.rows[idx]._tr, [
                idx, task_name, sample_sn, result_text, conclusion,
            ], center_cols={0, 3, 4})

    # ── 样品列表表格 ──
    if samples:
        doc.add_heading("样品台账", level=2)
        sample_headers = ["#", "SN", "批次号", "规格型号", "状态"]
        sample_table = doc.add_table(
            rows=1 + len(samples), cols=len(sample_headers), style="Table Grid"
        )
        sample_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _fill_row_cells(sample_table.rows[0]._tr, sample_headers, bold=True,
                        color="FFFFFF", shade="4F81BD", center_cols=set(range(len(sample_headers))))
        for idx, s in enumerate(samples, 1):
            _fill_row_cells(sample_table.rows[idx]._tr, [
                idx, s.sn, s.batch_no, s.spec or "",
                STATUS_MAP.get(s.status, s.status),
            ], center_cols={0})

    # ── 保存 ──
    out = filepath or str(output_dir / f"测试报告_{plan.name}_{datetime.now():%Y%m%d_%H%M}.docx")
    _validate_output_path(out, output_dir)
    try:
        doc.save(out)
    except (OSError, PermissionError) as e:
        logger.error("Word save failed: %s → %s", out, e)
        raise
    return os.path.abspath(out)


def export_dvpr_excel(
    output_dir: Path,
    plan: TestPlan,
    tasks: list[TestTask],
    results: list[TestResult],
    issues: list[Issue],
    samples: list[Sample],
    filepath: str | None = None,
) -> str:
    """导出 DVP&R 为 Excel (.xlsx)。"""
    from openpyxl import Workbook
    from openpyxl.styles import PatternFill, Font, Alignment

    s = excel_styles("2B579A")
    wb = Workbook()

    # ── Sheet 1: 概览 ──
    ws = wb.active
    ws.title = "概览"

    total = len(tasks)
    completed = sum(1 for t in tasks if t.status == "completed")
    total_pass = sum(1 for r in results if r.result == "pass")
    total_fail = sum(1 for r in results if r.result == "fail")
    total_results = len(results)
    pass_rate = f"{total_pass / total_results * 100:.1f}%" if total_results else "—"

    ws.merge_cells("A1:F1")
    ws["A1"].value = f"DVP&R — {plan.name}"
    ws["A1"].font = s["title_font"]("2B579A")
    ws["A1"].alignment = Alignment(horizontal="center")
    ws.row_dimensions[1].height = 30

    ws.merge_cells("A2:F2")
    ws["A2"].value = f"测试标准: {plan.test_standard or '—'}  |  导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    ws["A2"].font = s["sub_font"]
    ws["A2"].alignment = Alignment(horizontal="center")

    stats = [
        ("总任务数", total, "已完成", completed),
        ("测试结果", total_results, "通过率", pass_rate),
        ("Issue 数", len(issues), "样品数", len(samples)),
    ]
    excel_write_headers(ws, 4, ["指标", "值", "指标", "值"], s)
    for ri, (k1, v1, k2, v2) in enumerate(stats, 5):
        excel_write_row(ws, ri, [k1, v1, k2, v2], s)

    # ── Sheet 2: DVP&R 矩阵 ──
    ws2 = wb.create_sheet("DVP&R 矩阵")
    sample_ids = sorted({r.sample_id for r in results if r.sample_id})
    sample_map = {s.id: s.sn for s in samples if s.id is not None}
    lookup: dict[tuple[int, int], str] = {}
    for r in results:
        if r.task_id and r.sample_id:
            lookup[(r.task_id, r.sample_id)] = r.result

    headers = ["#", "测试项", "判定准则", "样品 SN"] + [
        sample_map.get(sid, f"#{sid}") for sid in sample_ids
    ] + ["结论"]
    excel_write_headers(ws2, 1, headers, s)

    _pass_fill = PatternFill(start_color="E8F5E9", end_color="E8F5E9", fill_type="solid")
    _fail_fill = PatternFill(start_color="FFEBEE", end_color="FFEBEE", fill_type="solid")

    for idx, task in enumerate(tasks, 2):
        row_data = [
            idx - 1,
            task.name or "",
            task.accept_criteria or "",
        ]
        task_sample_ids = sorted({r.sample_id for r in results if r.task_id == task.id and r.sample_id})
        task_sns = ", ".join(sample_map.get(sid, f"#{sid}") for sid in task_sample_ids)
        row_data.append(task_sns or "—")
        task_pass = task_fail = task_conditional = 0
        for sid in sample_ids:
            res = lookup.get((task.id, sid), "")
            if res == "pass":
                row_data.append("P"); task_pass += 1
            elif res == "fail":
                row_data.append("F"); task_fail += 1
            elif res == "conditional":
                row_data.append("C"); task_conditional += 1
            elif res == "skip":
                row_data.append("S")
            else:
                row_data.append("—")
        row_data.append(_judge_conclusion(task_pass, task_fail, task_conditional, 0,
                                          accept_criteria=task.accept_criteria or ""))
        excel_write_row(ws2, idx, row_data, s)

        for col_offset, sid in enumerate(sample_ids):
            res = lookup.get((task.id, sid), "")
            col_num = col_offset + 5
            cell = ws2.cell(row=idx, column=col_num)
            if res == "fail":
                cell.fill = _fail_fill
            elif res == "pass":
                cell.fill = _pass_fill

    from openpyxl.utils import get_column_letter
    ws2.column_dimensions["A"].width = 5
    ws2.column_dimensions["B"].width = 25
    ws2.column_dimensions["C"].width = 20
    ws2.column_dimensions["D"].width = 20
    for i in range(len(sample_ids)):
        ws2.column_dimensions[get_column_letter(5 + i)].width = 10
    ws2.column_dimensions[get_column_letter(5 + len(sample_ids))].width = 10

    # ── Sheet 3: Issue 汇总 ──
    if issues:
        ws3 = wb.create_sheet("Issue 汇总")
        issue_headers = ["ID", "Issue描述", "严重度", "状态", "DRI", "报告人", "解决结果"]
        excel_write_headers(ws3, 1, issue_headers, s)
        for ri, issue in enumerate(issues, 2):
            resolution_text = RESOLUTION_LABELS.get(issue.resolution, issue.resolution) or ""
            excel_write_row(ws3, ri, [
                issue.id, (issue.title or "")[:30], issue.severity,
                STATUS_MAP.get(issue.status, issue.status),
                issue.dri_name or "",
                issue.reporter_name or "",
                resolution_text,
            ], s)
        ws3.column_dimensions["A"].width = 6
        ws3.column_dimensions["B"].width = 30

    # ── Sheet 4: 签字栏 ──
    ws4 = wb.create_sheet("签字栏")
    ws4.merge_cells("A1:C1")
    ws4["A1"].value = "签字确认"
    ws4["A1"].font = s["title_font"]("2B579A")
    ws4["A1"].alignment = Alignment(horizontal="center")
    excel_write_headers(ws4, 3, ["编制", "审核", "批准"], s)
    for ri in range(4, 10):
        for ci in range(1, 4):
            cell = ws4.cell(row=ri, column=ci, value="")
            cell.border = s["thin_border"]
    ws4.column_dimensions["A"].width = 25
    ws4.column_dimensions["B"].width = 25
    ws4.column_dimensions["C"].width = 25

    return excel_save(wb, filepath, f"DVP&R_{plan.name}_{datetime.now():%Y%m%d_%H%M}.xlsx", output_dir)


def export_dvpr_docx(
    output_dir: Path,
    plan: TestPlan,
    tasks: list[TestTask],
    results: list[TestResult],
    issues: list[Issue],
    samples: list[Sample],
    filepath: str | None = None,
) -> str:
    """导出 DVP&R 为 Word (.docx)。"""
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    _f = _get_cjk_font()
    _BLUE = RGBColor(0x1E, 0x66, 0xA5)
    _GRAY = RGBColor(0x99, 0x99, 0x99)
    _DARK = RGBColor(0x33, 0x33, 0x33)
    _WHITE = RGBColor(0xFF, 0xFF, 0xFF)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Microsoft YaHei"
    font.size = Pt(9)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("DVP&R Report")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = _BLUE

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(f"Design Verification Plan & Report — {plan.name}")
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = _GRAY

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(f"测试标准: {plan.test_standard or '—'}  |  版本: V1.0  |  日期: {datetime.now().strftime('%Y-%m-%d')}")
    date_run.font.size = Pt(12)
    date_run.font.color.rgb = _GRAY

    # ── 概览 ──
    doc.add_heading("概览统计", level=2)
    total = len(tasks)
    completed = sum(1 for t in tasks if t.status == "completed")
    total_pass = sum(1 for r in results if r.result == "pass")
    total_fail = sum(1 for r in results if r.result == "fail")
    total_results = len(results)
    pass_rate = f"{total_pass / total_results * 100:.1f}%" if total_results else "—"

    for line_text in [
        f"总任务数: {total}  |  已完成: {completed}",
        f"测试结果: {total_results}  |  通过: {total_pass}  |  失败: {total_fail}  |  通过率: {pass_rate}",
        f"Issue 数: {len(issues)}  |  样品数: {len(samples)}",
    ]:
        p = doc.add_paragraph(line_text)
        p.paragraph_format.space_after = Pt(2)

    # ── DVP&R 矩阵 ──
    doc.add_heading("DVP&R 矩阵", level=2)
    sample_ids = sorted({r.sample_id for r in results if r.sample_id})
    sample_map = {s.id: s.sn for s in samples if s.id is not None}
    lookup: dict[tuple[int, int], str] = {}
    for r in results:
        if r.task_id and r.sample_id:
            lookup[(r.task_id, r.sample_id)] = r.result

    col_count = 4 + len(sample_ids) + 1
    dvpr_headers = ["#", "测试项", "判定准则", "样品 SN"] + [
        sample_map.get(sid, f"#{sid}") for sid in sample_ids
    ] + ["结论"]
    dvpr_table = doc.add_table(rows=1 + len(tasks), cols=col_count, style="Table Grid")
    dvpr_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    _fill_row_cells(dvpr_table.rows[0]._tr, dvpr_headers, bold=True,
                    color="FFFFFF", shade="2B579A", center_cols=set(range(len(dvpr_headers))))

    for idx, task in enumerate(tasks, 1):
        task_sample_ids = sorted({r.sample_id for r in results if r.task_id == task.id and r.sample_id})
        task_sns = ", ".join(sample_map.get(sid, f"#{sid}") for sid in task_sample_ids)
        row_vals = [str(idx), task.name or "", task.accept_criteria or "", task_sns or "—"]
        task_pass = task_fail = task_conditional = 0
        for sid in sample_ids:
            res = lookup.get((task.id, sid), "")
            if res == "pass":
                row_vals.append("P"); task_pass += 1
            elif res == "fail":
                row_vals.append("F"); task_fail += 1
            elif res == "conditional":
                row_vals.append("C"); task_conditional += 1
            elif res == "skip":
                row_vals.append("S")
            else:
                row_vals.append("—")
        row_vals.append(_judge_conclusion(task_pass, task_fail, task_conditional, 0,
                                          accept_criteria=task.accept_criteria or ""))

        tr = dvpr_table.rows[idx]._tr
        tc_elems = tr.findall(qn("w:tc"))
        for j, val in enumerate(row_vals):
            if j < len(tc_elems):
                shade_cell = None
                if 4 <= j < 4 + len(sample_ids):
                    sid = sample_ids[j - 4]
                    res = lookup.get((task.id, sid), "")
                    if res == "fail":
                        shade_cell = "FFEBEE"
                    elif res == "pass":
                        shade_cell = "E8F5E9"
                _fill_cell(tc_elems[j], str(val), shade=shade_cell, align="center")

    # ── Issue 汇总 ──
    if issues:
        doc.add_heading("Issue 追踪汇总", level=2)
        issue_headers = ["ID", "Issue描述", "严重度", "状态", "DRI", "报告人", "解决结果"]
        issue_table = doc.add_table(rows=1 + len(issues), cols=7, style="Table Grid")
        issue_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        _fill_row_cells(issue_table.rows[0]._tr, issue_headers, bold=True,
                        color="FFFFFF", shade="C0504D", center_cols=set(range(7)))
        for i, issue in enumerate(issues, 1):
            resolution_text = RESOLUTION_LABELS.get(issue.resolution, issue.resolution) or ""
            _fill_row_cells(issue_table.rows[i]._tr, [
                str(issue.id), (issue.title or "")[:30], issue.severity,
                STATUS_MAP.get(issue.status, issue.status),
                issue.dri_name or "",
                issue.reporter_name or "",
                resolution_text,
            ], center_cols={0})

    # ── 签字栏 ──
    doc.add_paragraph()
    sig_roles = ["编制 (Prepared)", "审核 (Reviewed)", "批准 (Approved)"]
    sig_table = doc.add_table(rows=1, cols=3)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, role in enumerate(sig_roles):
        tc = sig_table.rows[0]._tr.findall(qn("w:tc"))[j]
        _fill_cell_local(tc, f"{role}\n\n\n________________________", align="center")

    # ── 页脚 ──
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("ReliaTrack — DVP&R Report")
    run.font.size = Pt(8)
    run.font.color.rgb = _GRAY

    out = filepath or str(output_dir / f"DVP&R_{plan.name}_{datetime.now():%Y%m%d_%H%M}.docx")
    _validate_output_path(out, output_dir)
    try:
        doc.save(out)
    except (OSError, PermissionError) as e:
        logger.error("Word save failed: %s → %s", out, e)
        raise
    return os.path.abspath(out)


# ---------------------------------------------------------------------------
# 8D Report Word 导出 — 辅助函数
# ---------------------------------------------------------------------------

def _build_d4_content_docx(issue, fa_records):
    """构建 D4 根因分析内容。"""
    parts = []
    if issue.root_cause:
        parts.append(f"根因: {issue.root_cause}")
    if issue.failure_mode:
        parts.append(f"失效模式: {issue.failure_mode}")
    if fa_records:
        parts.append("\nFA 分析记录摘要:")
        for rec in fa_records:
            confirmed_labels = {0: "待定", 1: "确认", 2: "排除"}
            status = confirmed_labels.get(rec.confirmed, "待定")
            parts.append(
                f"  Step {rec.step_no}: {rec.step_title or ''}"
                f" — {rec.method or ''}"
                f" — 发现: {rec.findings or '—'}"
                f" — 可能原因: {rec.possible_cause or '—'}"
                f" [{status}]"
            )
    return "\n".join(parts) if parts else ""


def _build_d6_content_docx(capa_records):
    """构建 D6 实施验证内容。"""
    if not capa_records:
        return ""
    parts = []
    for rec in capa_records:
        status_labels = {
            "pending": "待执行", "in_progress": "进行中",
            "completed": "已完成", "verified": "已验证",
        }
        status = status_labels.get(rec.status, rec.status)
        parts.append(f"措施: {rec.action}")
        assignee_name = getattr(rec, 'assignee_name', '') or ''
        if assignee_name:
            parts.append(f"  负责人: {assignee_name}")
        parts.append(f"  状态: {status}")
        if rec.verification_result:
            parts.append(f"  验证结果: {rec.verification_result}")
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# 8D Report Word 导出
# ---------------------------------------------------------------------------

def export_8d_docx(
    output_dir: Path,
    issue: Issue,
    fa_records: list[FARecord] | None = None,
    capa_records: list[CAPARecord] | None = None,
    technician_name: str = "",
    task: TestTask | None = None,
    sample_sn: str = "",
    filepath: str | None = None,
) -> str:
    """导出 8D Problem Solving Report 为 Word (.docx)。

    结构与 export_8d_pdf 一致：基本信息表、D1-D8 八个章节、底部签字栏。
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    _BLUE = RGBColor(0x1E, 0x66, 0xA5)
    _GRAY = RGBColor(0x99, 0x99, 0x99)
    _DARK = RGBColor(0x33, 0x33, 0x33)
    _WHITE = RGBColor(0xFF, 0xFF, 0xFF)

    output_dir.mkdir(parents=True, exist_ok=True)
    out = filepath or str(
        output_dir / f"8D_Report_Issue{issue.id}_{datetime.now():%Y%m%d_%H%M}.docx"
    )

    doc = Document()

    # ── 页面设置 A4 ──
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    # ── 默认字体 ──
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Microsoft YaHei"
    font.size = Pt(9)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    # ── 标题 ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("8D Problem Solving Report")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = _BLUE

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"Issue #{issue.id} — {issue.title}")
    run.font.size = Pt(12)
    run.font.color.rgb = _GRAY

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run(f"报告日期: {datetime.now().strftime('%Y-%m-%d')}")
    run.font.size = Pt(12)
    run.font.color.rgb = _GRAY

    # ── 基本信息表 ──
    severity_labels = {
        "critical": "Critical (致命)",
        "major": "Major (严重)",
        "minor": "Minor (一般)",
        "cosmetic": "Cosmetic (外观)",
    }
    sev_text = severity_labels.get(issue.severity, issue.severity)
    status_text = STATUS_MAP.get(issue.status, issue.status)

    info = doc.add_table(rows=5, cols=4)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ["Issue 编号", str(issue.id), "严重度", sev_text],
        ["Issue描述", issue.title, "状态", status_text],
        ["报告日期", datetime.now().strftime("%Y-%m-%d"), "优先级", str(issue.priority)],
        ["DRI", issue.dri_name or "", "报告人", issue.reporter_name or ""],
        ["失效模式", issue.failure_mode or "", "解决结果", RESOLUTION_LABELS.get(issue.resolution, issue.resolution) or ""],
    ]
    for ri, row_data in enumerate(info_data):
        for ci, cell_text in enumerate(row_data):
            cell = info.cell(ri, ci)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.size = Pt(9)
            if ci in (0, 2):
                run.bold = True
                run.font.color.rgb = _BLUE
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # 蓝色背景通过 shading
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "E8F0FE")
                shading.set(qn("w:val"), "clear")
                cell._tc.get_or_add_tcPr().append(shading)
            else:
                run.font.color.rgb = _DARK
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 设置表格外框
    _set_table_border(info)

    # ── D1-D8 章节 ──
    # D1: 自动填充团队信息
    team_lines = []
    if technician_name:
        team_lines.append(f"负责人: {technician_name}")
    d1_content = "\n".join(team_lines) if team_lines else "(手写区)"

    # D2: 问题描述 — 补充测试条件和样品信息
    d2_parts = []
    if issue.description:
        d2_parts.append(issue.description)
    if task:
        tc_parts = []
        if task.name:
            tc_parts.append(f"测试项: {task.name}")
        if task.temperature:
            tc_parts.append(f"温度: {task.temperature}")
        if task.humidity:
            tc_parts.append(f"湿度: {task.humidity}")
        if task.accept_criteria:
            tc_parts.append(f"判定准则: {task.accept_criteria}")
        if tc_parts:
            d2_parts.append("测试条件: " + " | ".join(tc_parts))
    if sample_sn:
        d2_parts.append(f"样品: {sample_sn}")
    if issue.failure_stage:
        d2_parts.append(f"失效阶段: {issue.failure_stage}")
    d2_content = "\n".join(d2_parts) if d2_parts else ""

    # D3: 临时遏制 — 列出 pending/in_progress 的 CAPA
    capa_pending = []
    if capa_records:
        for _c in capa_records:
            if _c.status in ("pending", "in_progress"):
                capa_pending.append(f"• {_c.action} (状态: {_c.status})")
    d3_content = "\n".join(capa_pending) if capa_pending else "(手写区)"

    # D7: 预防再发 — 提示知识库
    d7_parts = ["(手写区 — 建议归档至知识库)"]
    if issue.failure_mode:
        d7_parts.append(f"失效模式: {issue.failure_mode} → 建议同步至知识库")
    d7_content = "\n".join(d7_parts)

    d_sections = [
        ("D1", "团队组建 (Establish the Team)", d1_content),
        ("D2", "问题描述 (Describe the Problem)", d2_content),
        ("D3", "临时遏制措施 (Interim Containment Actions)", d3_content),
        ("D4", "根因分析 (Root Cause Analysis)", _build_d4_content_docx(issue, fa_records)),
        ("D5", "纠正措施 (Corrective Actions)", RESOLUTION_LABELS.get(issue.resolution, issue.resolution) or ""),
        ("D6", "实施验证 (Implement & Validate)", _build_d6_content_docx(capa_records)),
        ("D7", "预防再发 (Prevent Recurrence)", "(手写区)"),
        ("D8", "结论与签字 (Congratulate the Team)", "(签字区)"),
    ]

    for d_label, d_title, d_content in d_sections:
        doc.add_paragraph()  # 间距

        # D 章节标题行 — 蓝底白字表格
        header_tbl = doc.add_table(rows=1, cols=2)
        header_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        # D 标签
        d_cell = header_tbl.cell(0, 0)
        d_cell.width = Cm(2.0)
        p = d_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(d_label)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = _WHITE

        # D 标题
        t_cell = header_tbl.cell(0, 1)
        p = t_cell.paragraphs[0]
        run = p.add_run(d_title)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = _WHITE

        # 蓝底
        for cell in [d_cell, t_cell]:
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "1E66A5")
            shading.set(qn("w:val"), "clear")
            cell._tc.get_or_add_tcPr().append(shading)

        # D 内容行
        content_tbl = doc.add_table(rows=1, cols=2)
        content_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        left_cell = content_tbl.cell(0, 0)
        left_cell.width = Cm(2.0)
        # 左侧留白
        right_cell = content_tbl.cell(0, 1)
        p = right_cell.paragraphs[0]
        content_text = d_content or ""
        if content_text.startswith("("):
            run = p.add_run(content_text)
            run.font.color.rgb = _GRAY
            run.font.size = Pt(9)
        else:
            for line_idx, line in enumerate(content_text.split("\n")):
                if line_idx > 0:
                    p = right_cell.add_paragraph()
                run = p.add_run(line)
                run.font.size = Pt(9)
                run.font.color.rgb = _DARK

        _set_table_border(content_tbl)
        _set_table_border(header_tbl)

    # ── 底部签字栏 ──
    doc.add_paragraph()
    sig_roles = ["编制 (Prepared)", "审核 (Reviewed)", "批准 (Approved)"]
    sig_tbl = doc.add_table(rows=2, cols=3)
    sig_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, role in enumerate(sig_roles):
        # 角色标签
        cell = sig_tbl.cell(0, ci)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(role)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = _DARK

        # 签字线
        cell2 = sig_tbl.cell(1, ci)
        p2 = cell2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run("________________________")
        run2.font.color.rgb = _GRAY
        run2.font.size = Pt(9)

    _set_table_border(sig_tbl)

    # ── 页脚 ──
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("ReliaTrack — 8D Problem Solving Report")
    run.font.size = Pt(8)
    run.font.color.rgb = _GRAY

    _validate_output_path(out, output_dir)
    try:
        doc.save(out)
    except (OSError, PermissionError) as e:
        logger.error("Word save failed: %s → %s", out, e)
        raise
    return os.path.abspath(out)